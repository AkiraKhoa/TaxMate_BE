import argparse
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "src/TaxMate.Infrastructure/Templates/Tax/2026/mau-01-tkn-cnkd.docx"


def cell(cell, text, bold=False, size=6):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text); r.bold = bold; r.font.name = "Times New Roman"; r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell_obj, fill="D9EAF7"):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill)
    cell_obj._tc.get_or_add_tcPr().append(shd)


def repeat(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def p(doc, text="", bold=False, size=9, center=False, italic=False, after=1):
    para = doc.add_paragraph(); para.paragraph_format.space_after = Pt(after)
    if center: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text); run.bold = bold; run.italic = italic
    run.font.name = "Times New Roman"; run.font.size = Pt(size)
    return para


def page(doc, printed):
    if printed != 11: doc.add_page_break()
    p(doc, str(printed), center=True, size=8, after=2)


def grid(doc, headers, rows, size=5.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers): cell(t.rows[0].cells[i], h, True, size); shade(t.rows[0].cells[i])
    repeat(t.rows[0])
    for ri, values in enumerate(rows, 1):
        for ci, value in enumerate(values): cell(t.rows[ri].cells[ci], value, False, size)
    return t


def grid_rows(doc, rows, size=5.5):
    t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, values in enumerate(rows):
        for ci, value in enumerate(values): cell(t.rows[ri].cells[ci], value, False, size)
    return t


def arow(stt, code, label, values=True):
    data = [stt, label, f"[{code}]", ""]
    data += [f"{{{{A{code}_{i}}}}}" if values else "0" for i in range(1, 16)]
    return data


def atotal_row():
    return ["6", "Tổng cộng", "[13]", ""] + [f"[13.{i}]\n{{{{A13_{i}}}}}" for i in range(1, 16)]


def aformula_row(stt, label, code, formula):
    return [stt, label, f"[{code}]", formula] + [""] * 15


doc = Document(); s = doc.sections[0]
s.orientation = WD_ORIENT.PORTRAIT; s.page_width = Cm(21); s.page_height = Cm(29.7)
s.top_margin = s.bottom_margin = Cm(1.05); s.left_margin = s.right_margin = Cm(1.0)
doc.styles["Normal"].font.name = "Times New Roman"; doc.styles["Normal"].font.size = Pt(9)

# Trang in 11 - header and taxpayer information.
page(doc, 11)
form_box = doc.add_table(rows=1, cols=2)
form_box.alignment = WD_TABLE_ALIGNMENT.CENTER
cell(form_box.cell(0, 0), "", size=7)
cell(form_box.cell(0, 1), "Mẫu số: 01/TKN-CNKD\n(Kèm theo Thông tư số 89/2026/TT-BTC ngày 30 tháng 6 năm 2026 của Bộ trưởng Bộ Tài chính)", True, 7)
p(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", True, 11, True, after=0)
p(doc, "Độc lập - Tự do - Hạnh phúc", True, 10, True, after=5)
p(doc, "THÔNG BÁO DOANH THU/TỜ KHAI THUẾ NĂM", True, 13, True, after=1)
p(doc, "(Áp dụng đối với hộ kinh doanh, cá nhân kinh doanh có doanh thu năm từ 01 tỷ đồng trở xuống; hộ kinh doanh, cá nhân kinh doanh nộp thuế TNCN theo phương pháp thuế suất nhân với doanh thu tính thuế đề nghị hoàn thuế; cá nhân hợp tác kinh doanh với tổ chức; cá nhân làm đại lý bán đúng giá đối với hoạt động bảo hiểm, xổ số, bán hàng đa cấp; cá nhân thực hiện hoạt động môi giới thuộc trường hợp phải nộp thuế mà chưa được tổ chức khấu trừ thuế, khai thuế thay, nộp thuế thay trong năm)", False, 7, True, True, 3)
p(doc, "{{AT_OR_BELOW_1B}} Hộ kinh doanh, cá nhân kinh doanh có doanh thu năm từ 01 tỷ đồng trở xuống")
p(doc, "{{NEW_BUSINESS}} Hộ kinh doanh, cá nhân kinh doanh mới ra kinh doanh có doanh thu năm từ 01 tỷ đồng trở xuống")
p(doc, "☐ Hộ kinh doanh, cá nhân kinh doanh nộp thuế TNCN theo phương pháp thuế suất nhân với doanh thu tính thuế đề nghị hoàn thuế")
p(doc, "☐ Cá nhân hợp tác kinh doanh với tổ chức; cá nhân làm đại lý bán đúng giá đối với hoạt động bảo hiểm, xổ số, bán hàng đa cấp; cá nhân thực hiện hoạt động môi giới thuộc trường hợp phải nộp thuế mà chưa được tổ chức khấu trừ thuế, khai thuế thay, nộp thuế thay trong năm")
p(doc, "☐ Cho phép điều chỉnh, bổ sung các tờ khai Mẫu số 01/CNKD đã kê khai theo Thông tư số 40/2021/TT-BTC, Thông tư số 18/2026/TT-BTC; tờ khai Mẫu số 02/CNKD-TMĐT đã kê khai theo Nghị định số 117/2025/NĐ-CP")
p(doc, "☐ Giảm thuế do thiên tai, dịch bệnh, hỏa hoạn, tai nạn và bệnh hiểm nghèo")
p(doc, "[01] Kỳ tính thuế:  [01a] Năm {{PERIOD_YEAR}} {{YEAR}}   [01b] 6 tháng đầu năm {{PERIOD_H1}} {{YEAR}}   [01c] 6 tháng cuối năm {{PERIOD_H2}} {{YEAR}}", True)
p(doc, "[02] Lần đầu: {{INITIAL}}          [03] Bổ sung lần thứ: {{SUPPLEMENT}}")
p(doc, "[04] Người nộp thuế: {{TAXPAYER_NAME}}", True)
p(doc, "[05] Mã số thuế: {{TAX_CODE}}", True)

# Trang in 12 - authorization and Section A header/[08]-[09].
page(doc, 12)
p(doc, "[06] Tổ chức/cá nhân kê khai, nộp thuế thay theo ủy quyền (nếu có): {{AUTHORIZED_NAME}}", True)
p(doc, "[06.1] Mã số thuế: {{AUTHORIZED_TAX_CODE}}")
p(doc, "[06.2] Văn bản ủy quyền (nếu có): Số................ ngày..... tháng..... năm.....")
p(doc, "[07] Tên tổ chức, cá nhân cung cấp dịch vụ làm thủ tục về thuế (nếu có): {{TAX_AGENT_NAME}}", True)
p(doc, "[07.1] Mã số thuế: {{TAX_AGENT_TAX_CODE}}")
p(doc, "[07.2] Hợp đồng dịch vụ làm thủ tục về thuế: Số {{TAX_AGENT_CONTRACT}} ngày {{TAX_AGENT_CONTRACT_DATE}}")
p(doc, "A. XÁC ĐỊNH NGHĨA VỤ THUẾ GTGT, TNCN", True, 10)
a_headers = ["STT", "Chỉ tiêu", "Mã chỉ tiêu", "Nhóm ngành nghề\n(1)", "Tổng doanh thu\n(2)", "DT không chịu GTGT\n(3)", "DT GTGT 0%\n(4)", "GTGT phải nộp\n(5)", "DT chịu TNCN\n(6)", "DT được trừ\n(7)", "TNCN phải nộp\n(8)", "GTGT trực tiếp nộp\n(9)", "TNCN trực tiếp nộp\n(10)", "GTGT khấu trừ/nộp thay\n(11)", "TNCN khấu trừ/nộp thay\n(12)", "GTGT phải nộp thêm\n(13)", "TNCN phải nộp thêm\n(14)", "GTGT nộp thừa\n(15)", "TNCN nộp thừa\n(16)"]
grid(doc, a_headers, [
    ["I", "Hoạt động sản xuất, kinh doanh hàng hóa, cung cấp dịch vụ (trừ hoạt động cho thuê bất động sản)"] + [""] * 17,
    arow("1", "08", "Hoạt động sản xuất, kinh doanh hàng hóa, cung cấp dịch vụ có địa điểm kinh doanh cố định"),
    arow("2", "09", "Hoạt động kinh doanh trên nền tảng thương mại điện tử, nền tảng số khác")], 4.0)

# Trang in 13 - remaining Section A and PIT formulas.
page(doc, 13)
grid_rows(doc, [
    arow("3", "10", "Hoạt động đại lý xổ số, bảo hiểm, bán hàng đa cấp", False),
    arow("4", "11", "Hoạt động môi giới", False),
    arow("5", "12", "Hoạt động cá nhân hợp tác kinh doanh với tổ chức", False),
    atotal_row(),
    aformula_row("7", "Số thuế TNCN từ hoạt động sản xuất, kinh doanh (trừ hoạt động cho thuê bất động sản) được giảm do thiên tai, dịch bệnh, hỏa hoạn, tai nạn và bệnh hiểm nghèo theo Quyết định giảm thuế của cơ quan thuế", "14", "0"),
    aformula_row("8", "Số thuế TNCN còn phải nộp", "15", "[15]=[13.13]-[14]>=0"),
    aformula_row("9", "Số thuế TNCN được miễn từ 50.000 đồng trở xuống", "16", "[16]=[15]<=50.000"),
    aformula_row("10", "Số thuế TNCN còn phải nộp sau miễn số thuế từ 50.000 đồng trở xuống", "17", "[17]=[15]-[16]"),
    aformula_row("11", "Số thuế TNCN nộp thừa", "18", "[18]=[13.13]-[14]<0 hoặc [13.15]+[14]>0"),
    ["II", "Hoạt động cho thuê bất động sản", "[19]", ""] + [f"[19.{i}]" for i in range(1, 16)],
    aformula_row("1", "Số thuế TNCN từ hoạt động cho thuê bất động sản", "20", "0")], 3.9)

# Trang in 14 - rental total and excise tax.
page(doc, 14)
grid(doc, ["STT", "Chỉ tiêu", "Mã chỉ tiêu", "Giá trị/công thức"], [
    ["2", "Số thuế TNCN nộp thừa", "[21]", "[21]=[19.15]+[20]"],
    ["B", "KÊ KHAI THUẾ TIÊU THỤ ĐẶC BIỆT", "", ""]], 7)
grid(doc, ["STT", "Hàng hóa, dịch vụ chịu thuế TTĐB", "Mã chỉ tiêu", "Đơn vị tính", "Doanh thu tính thuế TTĐB", "Thuế suất", "Số thuế phải nộp"], [
    ["I", "Trụ sở kinh doanh", "", "", "", "", ""],
    ["1", "Hàng hóa, dịch vụ A", "[22a]", "", "", "", ""],
    ["2", "Hàng hóa, dịch vụ B", "[22b]", "", "", "", ""],
    ["II", "Mã địa điểm kinh doanh 1 / Tên địa điểm kinh doanh 1", "", "", "", "", ""],
    ["III", "Tổng cộng", "[23]", "", "[23.1]", "", "[23.2]"],
    ["IV", "Số thuế được miễn từ 50.000 đồng trở xuống", "[24]", "", "[24]=[23.2]<=50.000", "", ""],
    ["V", "Số thuế còn phải nộp", "[25]", "", "[25]=[23.2]-[24]", "", ""]], 6)

# Trang in 15 - resource/environmental table opening.
page(doc, 15)
p(doc, "C. KÊ KHAI THUẾ TÀI NGUYÊN, BẢO VỆ MÔI TRƯỜNG, PHÍ BẢO VỆ MÔI TRƯỜNG", True, 10)
c_headers = ["STT", "Tài nguyên, hàng hóa, sản phẩm", "Trụ sở/Mã địa điểm", "Tên địa điểm", "Mã chỉ tiêu", "Đơn vị tính", "Sản lượng/Số lượng", "Giá tính thuế/mức thuế phí", "Thuế suất/Hệ số K", "Số thuế/phí phải nộp"]
grid(doc, c_headers, [
    ["I", "Khai thuế tài nguyên", "", "", "", "", "", "", "", ""],
    ["1", "Tài nguyên C", "", "", "[26a]", "", "", "", "", "(10)=(7)*(8)*(9)"],
    ["2", "Tài nguyên D", "", "", "[26b]", "", "", "", "", ""],
    ["", "Tổng cộng", "", "", "", "", "", "", "", "[27]"],
    ["", "Số thuế được miễn từ 50.000 đồng trở xuống", "", "", "", "", "", "", "", "[28]=[27]<=50.000"],
    ["", "Số thuế còn phải nộp", "", "", "", "", "", "", "", "[29]=[27]-[28]"],
    ["II", "Khai thuế bảo vệ môi trường", "", "", "", "", "", "", "", ""]], 5.5)

# Trang in 16 - environmental/mineral continuation and Section D heading.
page(doc, 16)
grid(doc, c_headers, [
    ["2.1", "Hàng hóa E", "", "", "[30a]", "", "", "", "", "(10)=(7)*(8)"],
    ["2.2", "Hàng hóa G", "", "", "[30b]", "", "", "", "", ""],
    ["", "Tổng cộng", "", "", "", "", "", "", "", "[31]"],
    ["", "Số thuế được miễn từ 50.000 đồng trở xuống", "", "", "", "", "", "", "", "[32]"],
    ["", "Số thuế còn phải nộp", "", "", "", "", "", "", "", "[33]=[31]-[32]"],
    ["III", "Khai phí bảo vệ môi trường đối với khai thác khoáng sản", "", "", "", "", "", "", "", ""],
    ["3.1", "Khoáng sản X", "", "", "[34a]", "", "", "", "", "(10)=(7)*(8)*(9)"],
    ["3.2", "Khoáng sản Y", "", "", "[34b]", "", "", "", "", ""],
    ["", "Tổng cộng", "", "", "", "", "", "", "", "[35]"],
    ["", "Số phí được miễn từ 50.000 đồng trở xuống", "", "", "", "", "", "", "", "[36]=[35]<=50.000"],
    ["", "Số phí còn phải nộp", "", "", "", "", "", "", "", "[37]=[35]-[36]"]], 5)
p(doc, "D. HỖ TRỢ THÔNG TIN NỘP THUẾ", True, 10)

# Trang in 17 - payment support and overpayment requests.
page(doc, 17)
grid(doc, ["STT [38]", "Mã địa điểm [39]", "Nội dung khoản nộp NSNN [40]", "Số tiền [41]", "Chương [42]", "Tiểu mục [43]", "Địa bàn hành chính [44]", "Cơ quan thu [45]", "Cơ quan thuế [46]", "Hạn nộp [47]"], [
    ["", "", "", "", "", "", "", "", "", ""],
    ["", "", "Tổng cộng", "[48]", "", "", "", "", "", ""]], 5.5)
p(doc, "E. ĐỀ NGHỊ XỬ LÝ KHOẢN NỘP THỪA", True, 10)
p(doc, "[49] Đề nghị hoàn trả: ....................")
p(doc, "[49.1] Số thuế GTGT: ....................")
p(doc, "[49.2] Số thuế TNCN từ hoạt động sản xuất kinh doanh (trừ hoạt động cho thuê nhà): ....................")
p(doc, "[49.3] Số thuế TNCN từ hoạt động cho thuê nhà: ....................")
p(doc, "[49.4] Tổng cộng: ....................")
p(doc, "[50] Số bù trừ cho các phát sinh của kỳ sau: ....................")
p(doc, "[50.1] Số thuế GTGT: ....................")
p(doc, "[50.2] Số thuế TNCN từ hoạt động sản xuất kinh doanh (trừ hoạt động cho thuê nhà): ....................")
p(doc, "[50.3] Số thuế TNCN từ hoạt động cho thuê nhà: ....................")
p(doc, "[50.4] Tổng cộng: ....................")
p(doc, "[51] Thông tin hoàn trả:")
p(doc, "[51.1] Tên chủ tài khoản: ....................  [51.2] Tài khoản số: ....................  [51.3] Tại Ngân hàng/KBNN: ....................")

# Trang in 18 - offset table, declaration/signature and official notes.
page(doc, 18)
p(doc, "Thông tin người nộp thuế đề nghị bù trừ khoản nộp thừa với khoản nợ, khoản thu phát sinh hoặc đề nghị hoàn kiêm bù trừ thu ngân sách nhà nước:", size=8)
grid(doc, ["STT [52]", "Mã số thuế [53]", "Tên NNT [54]", "Mã định danh khoản phải nộp [55]", "Nội dung khoản nợ/phát sinh [56]", "Chương [57]", "Tiểu mục [58]", "Cơ quan thu [59]", "Địa bàn hành chính [60]", "Hạn nộp [61]", "Số tiền còn phải nộp [62]", "Số tiền đề nghị bù trừ [63]", "Số tiền còn phải nộp sau bù trừ [64]=[62]-[63]"], [
    ["", "", "", "", "", "", "", "", "", "", "", "", ""],
    ["", "", "", "", "", "", "", "", "", "", "", "", ""]], 4.7)
p(doc, "Tôi cam đoan những nội dung kê khai trên là đúng và chịu trách nhiệm trước pháp luật về những nội dung đã khai.", size=8)
sig = grid(doc, ["NGƯỜI TRỰC TIẾP THỰC HIỆN DỊCH VỤ LÀM THỦ TỤC VỀ THUẾ", "{{DECLARATION_DATE}}\nNGƯỜI NỘP THUẾ hoặc ĐẠI DIỆN HỢP PHÁP CỦA NGƯỜI NỘP THUẾ"], [["Họ và tên: ............................\nChứng chỉ nghiệp vụ chuyên môn về thuế số: ............", "(Chữ ký, ghi rõ họ tên; chức vụ và đóng dấu (nếu có)/Xác nhận điện tử/Ký điện tử)"]], 7)
p(doc, "Ghi chú:", True, 8)
p(doc, "- Đối với trường hợp hộ kinh doanh, cá nhân kinh doanh có doanh thu năm từ 01 tỷ đồng trở xuống thì chỉ thực hiện thông báo doanh thu, không thực hiện khai số thuế GTGT, thuế TNCN phải nộp.", False, 7)
p(doc, "- Chỉ tiêu (1) Nhóm ngành nghề tại mục A khai như sau: (1) Phân phối, cung cấp hàng hóa; (2) Dịch vụ, xây dựng không bao thầu nguyên vật liệu; (3) Hoạt động cho thuê tài sản trừ bất động sản; (4) Sản xuất, vận tải, dịch vụ có gắn với hàng hóa, xây dựng có bao thầu nguyên vật liệu; (5) Hoạt động kinh doanh khác.", False, 7)

OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT); print(OUT)

parser = argparse.ArgumentParser(); parser.add_argument("--qa-output"); args = parser.parse_args()
if args.qa_output:
    qa = Document(OUT)
    values = {"{{AT_OR_BELOW_1B}}":"☒","{{NEW_BUSINESS}}":"☐","{{YEAR}}":"2026",
        "{{PERIOD_YEAR}}":"☒","{{PERIOD_H1}}":"☐","{{PERIOD_H2}}":"☐","{{INITIAL}}":"☒",
        "{{SUPPLEMENT}}":"","{{TAXPAYER_NAME}}":"Nguyễn Văn Minh","{{TAX_CODE}}":"0312345678",
        "{{AUTHORIZED_NAME}}":"","{{AUTHORIZED_TAX_CODE}}":"","{{TAX_AGENT_NAME}}":"",
        "{{TAX_AGENT_TAX_CODE}}":"","{{TAX_AGENT_CONTRACT}}":"","{{TAX_AGENT_CONTRACT_DATE}}":"",
        "{{DECLARATION_DATE}}":"26/08/2026"}
    for code in ["08","09","10","11","12","13"]:
        for i in range(1,16): values[f"{{{{A{code}_{i}}}}}"] = "0"
    values["{{A08_1}}"] = values["{{A13_1}}"] = "900.000.000"
    values["{{A08_2}}"] = values["{{A13_2}}"] = "900.000.000"
    for para in list(qa.paragraphs)+[x for t in qa.tables for row in t.rows for c in row.cells for x in c.paragraphs]:
        for run in para.runs:
            for marker, value in values.items():
                if marker in run.text: run.text = run.text.replace(marker, value)
    q = Path(args.qa_output); q.parent.mkdir(parents=True, exist_ok=True); qa.save(q); print(q)
